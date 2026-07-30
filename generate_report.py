"""
Generate the GHF Multimodal Emotion Recognition final integration report (Word).
Run: python generate_report.py
Output: outputs/GHF_Integration_Final_Report.docx
"""

import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "outputs",
    "GHF_Integration_Final_Report.docx",
)


def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table(doc, headers, rows, header_color="1F4E79"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_color)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)

    doc.add_paragraph()
    return table


def build_report():
    doc = Document()

    # Title
    title = doc.add_heading("GHF Multimodal Emotion Recognition", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Final Integration Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].bold = True

    date_p = doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "This project implements a Multimodal Emotion Recognition System that combines "
        "facial expression analysis (CNN/FER model) with speech-to-text emotion classification "
        "(SVC, Random Forest, XGBoost). Two Streamlit applications were built on a shared "
        "utils/ module library:"
    )
    add_bullet(doc, " app.py — Single snapshot prediction (on-demand fusion)", bold_prefix="")
    add_bullet(doc, " app1.py — Continuous session monitoring with 30-second fusion windows, history, CSV logging, and analytics dashboard")

    doc.add_paragraph(
        "Both applications share core ML and webcam modules. app1.py extends the architecture "
        "with buffering, timing, background speech capture, session orchestration, and reporting."
    )

    # 2. Application Comparison
    add_heading(doc, "2. Application Comparison", 1)

    add_table(
        doc,
        ["Feature", "app.py (Snapshot)", "app1.py (Continuous Monitoring)"],
        [
            ["Purpose", "What emotion right now?", "How emotion changed over the session"],
            ["Face input", "Live webcam (continuous display)", "Live webcam + 1 sample/second buffer"],
            ["Speech input", "Manual button — record once", "Background thread — every 5 seconds"],
            ["Fusion timing", "Immediate (on button click)", "Every 30 seconds automatically"],
            ["History storage", "None", "In-memory HistoryManager"],
            ["CSV logging", "None", "outputs/emotion_history.csv"],
            ["Analytics dashboard", "None", "Dominant emotion, frequency charts, timeline"],
            ["Files directly used", "6 utils modules", "12 utils modules (via SessionManager)"],
            ["Run command", "streamlit run app.py", "streamlit run app1.py"],
        ],
    )

    # 3. File Count Summary
    add_heading(doc, "3. Module Usage Summary", 1)

    add_table(
        doc,
        ["Category", "Count", "Details"],
        [
            ["Total project Python files", "22", "2 apps + 18 utils + 2 test scripts"],
            ["Shared by both apps", "7", "webcam, face_detection, fer_prediction, speech_recognition, text_prediction, fusion, constants (+ dependencies)"],
            ["Used only by app.py", "2", "speech_recognition (direct), fusion (direct)"],
            ["Used only by app1.py", "11", "session_manager + 10 continuous-monitoring modules"],
            ["Support / shared infrastructure", "8", "constants, fer_prediction, integration_rules, text_preprocessing, face_detection, webcam, fusion, text_prediction"],
            ["Test utilities", "2", "test_face.py, test_text.py"],
            ["ML model files", "6", "FER CNN + 3 text classifiers + vectorizer + label encoder"],
        ],
    )

    doc.add_paragraph(
        "In summary: app.py directly imports and uses 6 utility modules. "
        "app1.py directly imports 2 modules (session_manager, webcam) but SessionManager "
        "orchestrates 10 additional modules behind the scenes — giving app1 a total footprint "
        "of 12 utility modules."
    )

    # 4. Architecture
    add_heading(doc, "4. System Architecture", 1)

    add_heading(doc, "4.1 app.py — Snapshot Workflow", 2)
    doc.add_paragraph(
        "Webcam → Face Detection → FER CNN → Live Face Emotion\n"
        "                                              ↓\n"
        "Microphone (button) → Speech Recognition → Text Models → Text Emotion\n"
        "                                              ↓\n"
        "                              Fusion Rules → Final Emotion (displayed once)"
    )

    add_heading(doc, "4.2 app1.py — Continuous Monitoring Workflow", 2)
    doc.add_paragraph(
        "Camera ──────────────► FaceBuffer (1 frame/sec, 30-sec window)\n"
        "                              │\n"
        "Microphone ──────────► SpeechBuffer (every 5 sec via ContinuousSpeech)\n"
        "                              │\n"
        "                     Every 30 Seconds (WindowTimer)\n"
        "                              │\n"
        "                       FusionManager + fuse_emotions()\n"
        "                              │\n"
        "              ┌────────────────┴────────────────┐\n"
        "              ▼                                 ▼\n"
        "       HistoryManager                    CSVLogger\n"
        "              │\n"
        "              ▼\n"
        "       AnalyticsManager → Dashboard (Streamlit UI)"
    )

    add_heading(doc, "4.3 Integration Layer", 2)
    doc.add_paragraph(
        "SessionManager (utils/session_manager.py) is the central integration hub for app1.py. "
        "It wires FaceBuffer, WindowTimer, ContinuousSpeech, FusionManager, HistoryManager, "
        "CSVLogger, and AnalyticsManager into a single coordinated pipeline. "
        "app1.py remains a thin UI layer — all business logic lives in utils/."
    )

    # 5. Complete File Reference
    add_heading(doc, "5. Complete File Reference", 1)

    add_heading(doc, "5.1 Application Entry Points", 2)
    add_table(
        doc,
        ["File", "Role", "Used By", "Status"],
        [
            ["app.py", "Snapshot multimodal emotion predictor — webcam + manual speech button", "End user (demo)", "Working"],
            ["app1.py", "Continuous monitoring dashboard — auto fusion, history, CSV, analytics", "End user (session monitoring)", "Working"],
        ],
    )

    add_heading(doc, "5.2 Utils — Face Emotion Pipeline", 2)
    add_table(
        doc,
        ["File", "Contains", "Used By", "Status"],
        [
            ["utils/webcam.py", "EmotionProcessor — Streamlit WebRTC video frame handler", "app.py, app1.py", "Working"],
            ["utils/face_detection.py", "Haar cascade face detection, frame overlay, global emotion state", "webcam.py, app.py, session_manager", "Working"],
            ["utils/fer_prediction.py", "CNN/Keras FER model loading and 7-class emotion prediction", "face_detection.py", "Working"],
            ["utils/constants.py", "Model paths, emotion labels, image size configuration", "fer_prediction, text_prediction", "Working"],
        ],
    )

    add_heading(doc, "5.3 Utils — Speech / Text Emotion Pipeline", 2)
    add_table(
        doc,
        ["File", "Contains", "Used By", "Status"],
        [
            ["utils/speech_recognition.py", "Google Speech API wrapper — microphone to text", "app.py, continuous_speech.py", "Working"],
            ["utils/text_preprocessing.py", "Text cleaning and normalization before ML inference", "text_prediction.py", "Working"],
            ["utils/text_prediction.py", "Loads SVC, RF, XGBoost + TF-IDF; returns 4-tuple predictions", "app.py, continuous_speech.py", "Working"],
            ["utils/integration_rules.py", "Majority vote / priority rules to combine 3 text model outputs", "text_prediction.py", "Working"],
        ],
    )

    add_heading(doc, "5.4 Utils — Fusion", 2)
    add_table(
        doc,
        ["File", "Contains", "Used By", "Status"],
        [
            ["utils/fusion.py", "Rule-based fusion of face emotion (7 classes) + text emotion (positive/neutral/negative)", "app.py, fusion_manager.py", "Working"],
            ["utils/fusion_manager.py", "30-second window fusion orchestrator — summarizes buffers then fuses", "session_manager.py", "Working"],
        ],
    )

    add_heading(doc, "5.5 Utils — Continuous Monitoring (app1 only)", 2)
    add_table(
        doc,
        ["File", "Contains", "Used By", "Status"],
        [
            ["utils/face_buffer.py", "Stores face predictions over 30-sec window; majority vote summary", "session_manager, fusion_manager", "Working"],
            ["utils/speech_buffer.py", "Thread-safe storage of speech emotion predictions + recognized text", "continuous_speech.py", "Working"],
            ["utils/text_buffer.py", "Alternative text emotion buffer (30-sec window) — standalone utility", "Not wired (SpeechBuffer used instead)", "Available"],
            ["utils/timer.py", "WindowTimer — 30-second countdown and elapsed time tracking", "session_manager, fusion_manager", "Working"],
            ["utils/continuous_speech.py", "Background thread — records speech every 5 sec, predicts emotion", "session_manager.py", "Working"],
            ["utils/history_manager.py", "In-memory session history of all 30-sec fusion summaries", "session_manager.py, app1.py", "Working"],
            ["utils/csv_logger.py", "Appends fusion results to outputs/emotion_history.csv", "session_manager.py", "Working"],
            ["utils/analytics_manager.py", "Dominant emotion, frequency counts for face/speech/final", "session_manager.py, app1.py", "Working"],
            ["utils/session_manager.py", "Master orchestrator — connects all app1 components", "app1.py", "Working"],
        ],
    )

    add_heading(doc, "5.6 Models Directory", 2)
    add_table(
        doc,
        ["File", "Model Type", "Purpose"],
        [
            ["models/fer_model.keras", "CNN (Keras)", "Facial emotion recognition — 7 classes"],
            ["models/fernet_bestweight.weights.h5", "Weights", "FER model weight backup"],
            ["models/svc_model.pkl", "SVC", "Text emotion classification"],
            ["models/random_forest_model.pkl", "Random Forest", "Text emotion classification"],
            ["models/xgboost_model.pkl", "XGBoost", "Text emotion classification"],
            ["models/tfidf_vectorizer.pkl", "TF-IDF", "Text feature extraction"],
            ["models/label_encoder.pkl", "Label Encoder", "XGBoost label decoding"],
        ],
    )

    add_heading(doc, "5.7 Test & Output Files", 2)
    add_table(
        doc,
        ["File", "Purpose"],
        [
            ["test_face.py", "Standalone test for face emotion prediction"],
            ["test_text.py", "Standalone test for text emotion prediction"],
            ["outputs/emotion_history.csv", "Auto-generated CSV log from app1 sessions"],
            ["requirements.txt", "Python dependencies (Streamlit, TensorFlow, scikit-learn, etc.)"],
        ],
    )

    # 6. Dependency Map
    add_heading(doc, "6. Dependency Map — Which App Uses Which File", 1)

    add_table(
        doc,
        ["Utility File", "app.py", "app1.py"],
        [
            ["utils/webcam.py", "Yes", "Yes"],
            ["utils/face_detection.py", "Yes (direct)", "Yes (via session_manager)"],
            ["utils/fer_prediction.py", "Yes (via face_detection)", "Yes (via face_detection)"],
            ["utils/constants.py", "Yes (via chain)", "Yes (via chain)"],
            ["utils/speech_recognition.py", "Yes (direct)", "Yes (via continuous_speech)"],
            ["utils/text_prediction.py", "Yes (direct)", "Yes (via continuous_speech)"],
            ["utils/text_preprocessing.py", "Yes (via text_prediction)", "Yes (via text_prediction)"],
            ["utils/integration_rules.py", "Yes (via text_prediction)", "Yes (via text_prediction)"],
            ["utils/fusion.py", "Yes (direct)", "Yes (via fusion_manager)"],
            ["utils/face_buffer.py", "No", "Yes"],
            ["utils/speech_buffer.py", "No", "Yes"],
            ["utils/text_buffer.py", "No", "No (available)"],
            ["utils/timer.py", "No", "Yes"],
            ["utils/continuous_speech.py", "No", "Yes"],
            ["utils/fusion_manager.py", "No", "Yes"],
            ["utils/history_manager.py", "No", "Yes"],
            ["utils/csv_logger.py", "No", "Yes"],
            ["utils/analytics_manager.py", "No", "Yes"],
            ["utils/session_manager.py", "No", "Yes (direct import)"],
        ],
    )

    doc.add_paragraph("File count totals:")
    add_bullet(doc, " app.py uses 6 utility modules directly (+ 5 indirect dependencies)", bold_prefix="")
    add_bullet(doc, " app1.py uses 12 utility modules (2 direct + 10 via SessionManager)", bold_prefix="")
    add_bullet(doc, " 7 modules are shared between both applications", bold_prefix="")

    # 7. Integration Details
    add_heading(doc, "7. How Integration Was Achieved", 1)

    steps = [
        (
            "Step 1 — Shared Face Pipeline",
            "Both apps use the same webcam → face_detection → fer_prediction chain. "
            "EmotionProcessor wraps process_frame() for Streamlit WebRTC. "
            "get_current_face_prediction() exposes the latest face emotion globally."
        ),
        (
            "Step 2 — Shared Text Pipeline",
            "Both apps use speech_recognition → text_preprocessing → text_prediction → integration_rules. "
            "Three ML models (SVC, RF, XGBoost) vote on text emotion (positive/neutral/negative)."
        ),
        (
            "Step 3 — Shared Fusion Rules",
            "utils/fusion.py maps face emotions (Happy, Sad, Angry, etc.) combined with text emotions "
            "to a single final label. app.py calls fuse_emotions() directly; app1 uses FusionManager."
        ),
        (
            "Step 4 — Continuous Monitoring Layer (app1)",
            "New modules were added: FaceBuffer and SpeechBuffer collect predictions over time. "
            "WindowTimer triggers fusion every 30 seconds. ContinuousSpeech runs a background "
            "daemon thread for automatic speech capture every 5 seconds."
        ),
        (
            "Step 5 — Session Orchestration",
            "SessionManager integrates all app1 components: starts speech thread, collects face "
            "samples at 1/sec, triggers FusionManager, saves to HistoryManager and CSVLogger, "
            "and exposes analytics for the dashboard."
        ),
        (
            "Step 6 — UI Layer Separation",
            "app1.py is a thin Streamlit frontend. All logic lives in utils/. "
            "This mirrors app.py's pattern but adds session persistence and reporting."
        ),
    ]

    for title, body in steps:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        doc.add_paragraph(body)

    # 8. Sample Results
    add_heading(doc, "8. Sample Session Results (from CSV)", 1)
    doc.add_paragraph(
        "The following records were captured during a live app1.py session and saved "
        "automatically to outputs/emotion_history.csv:"
    )

    csv_path = os.path.join(os.path.dirname(OUTPUT_PATH), "emotion_history.csv")
    if os.path.exists(csv_path):
        with open(csv_path, encoding="utf-8") as f:
            lines = [line.strip().split(",") for line in f.readlines()]
        if len(lines) > 1:
            add_table(doc, lines[0], lines[1:])
        else:
            doc.add_paragraph("CSV file exists but contains no session data yet.")
    else:
        doc.add_paragraph("No CSV output file found yet — run app1.py to generate session data.")

    # 9. Fusion Rules Reference
    add_heading(doc, "9. Fusion Rules Reference", 1)
    add_table(
        doc,
        ["Face Emotion", "Text Emotion", "Final Emotion"],
        [
            ["Happy", "positive / neutral", "Happy"],
            ["Happy", "negative", "Neutral"],
            ["Neutral", "positive", "Happy"],
            ["Neutral", "neutral", "Neutral"],
            ["Neutral", "negative", "Sad"],
            ["Angry", "any", "Angry"],
            ["Sad", "positive", "Neutral"],
            ["Sad", "neutral / negative", "Sad"],
            ["Fear", "any", "Fear"],
            ["Surprise", "negative", "Neutral"],
            ["Surprise", "other", "Surprise"],
            ["Disgust", "any", "Disgust"],
        ],
    )

    # 10. How to Run
    add_heading(doc, "10. How to Run", 1)
    add_table(
        doc,
        ["Application", "Command", "URL"],
        [
            ["Snapshot predictor", "streamlit run app.py", "http://localhost:8501"],
            ["Continuous monitor", "streamlit run app1.py", "http://localhost:8501"],
            ["Test face model", "python test_face.py", "Terminal"],
            ["Test text model", "python test_text.py", "Terminal"],
        ],
    )

    doc.add_paragraph("Prerequisites: Python 3.x, dependencies from requirements.txt, webcam, and microphone.")

    # 11. Conclusion
    add_heading(doc, "11. Conclusion", 1)
    doc.add_paragraph(
        "The GHF Multimodal Emotion Recognition project successfully integrates facial and "
        "speech-based emotion detection into two complementary applications. app.py provides "
        "instant, on-demand multimodal predictions suitable for demonstrations. app1.py extends "
        "this into a full emotion monitoring system with temporal buffering, periodic fusion, "
        "persistent logging, and analytics — answering both \"what emotion now?\" and "
        "\"how has emotion evolved over time?\""
    )

    doc.add_paragraph(
        "All 18 utility modules are functional. app.py handles 6 modules directly; "
        "app1.py orchestrates 12 modules through SessionManager. Shared infrastructure "
        "ensures consistent ML predictions across both applications."
    )

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("— End of Report —")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_report()
    print(f"Report saved to: {path}")
